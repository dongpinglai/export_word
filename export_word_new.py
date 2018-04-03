#!/usr/bin/env python
# -*- encoding: utf-8 -*-

import docx

class Text(object):
    """文本对象"""
    def __init__(self, text):
        self.text= text


class Image(object):
    """图片对象"""
    def __init__(self, title, identity, text):
        self.title = title
        self.identity = identity
        self.text = text

class Table(object):
    """表格对象"""
    def __init__(self, title, datas, table_type):
        self.title = title
        self.datas = datas
        self.table_type = table_type


class Parts(object):
    """
    不同集合部分的父类
    """
    _TYPE = None
    def __init__(self, iterable):
        if _TYPE is None:
            raise ValueError("_TYPE is None")
        self._items = []
        for item in iterable:
            if _TYPE in ["image"]:
                item = Image(**item)
            elif _TYPE in ["table"]:
                item = Table(**item)
            elif _TYPE in ["text"]:
                item = Text(item)
            else:
                raise ValueError("_TYPE can't be %s" % _TYPE)
            self._items.append(item)

    def __getitem__(self, index):
        """重写self[i]"""
        return self._items[index]

    def __getslice__(self, start=None, end=None):
        """重写self[i:j]"""
        return self._items[start: end]

    def __getattr__(self, attr):
        """重写self.attr"""
        _rel_in_line = False # 判断元素之间的关系，是否在同一行
        # 获取所有元素
        if attr == "all":
            return (self[:], _rel_in_line)
        # 获取单个元素
        elif "-" not in attr or ":" not in attr:
            return ([self[int(attr)]], _rel_in_line)
        else:
            # 获取多个元素
            if "-" in attr: 
                start, end = attr.split("-")
            if ":" in attr: 
                start, end = attr.split(":")
                _rel_in_line = True
            if start == "":
                start = None
            if end == "":
                end = None
            if start:
                start = int(start)
            if end:
                end = int(end)
            return (self[start, end], _rel_in_line)

    def __iter__(self):
        return self._items

        
class Images(Parts):
    """图片集合"""
    _TYPE = "image"

class Texts(Parts):
    """文本集合"""
    _TYPE = "text"

class Tables(Parts):
    """表格集合"""
    _TYPE = "table"
    

class Section(object):
    """区域对象"""
    def __init__(self, title, images, texts, tables, file_manager, sequence=[]):
        """
        title: 标题
        images: 图片
        texts: 文字
        tables: 表格
        sequence: 标题，图片， 文字，表格，的排列顺序.eg:
            ["title",
            "images.1-3",
            "texts",
            "tables.0",
            "tables.1:2",
            "tables.1:"
            ]
        """
        self.title = Texts([title])
        self.images = Images(images)
        self.texts = Texts(texts)
        self.tables = Tables(tables)
        self.file_manager = file_manager
        self.sequence = sequence
        self._render_items = []

    def gen_render_items(self):
        """得到要渲染的真实数据"""
        for it in sequence:
            attr_name, pos = self._get_attr_pos(it)
            attr = getattr(self, attr_name)
            item = getattr(attr, pos)
            self._render_items.append(item)

    def _get_attr_pos(self, seq_item):
        """
        根据sequence中的元素，获取属性对象名和属性对象中真实数据的位置信息
        """
        if "." in seq_item:
            attr_pos = seq_item.split(".")
        else:
            attr_pos = seq_item
        if len(attr_pos) == 1:
            attr = attr_pos[0]
            pos = "all"
        elif len(attr_pos) > 1:
            attr, pos = attr_pos
        return attr, pos
            
    def render(self, doc):
        """渲染区域
        doc: docx.Document对象
        """
        render_items = self._render_items
        for r_item in render_items:
            real_items, rel_in_line = r_item
            if rel_in_inline:
                self._render_inline(real_items, doc)
            else:
                self._render_not_inline(real_items, doc)
    
    def _render_inline(self, real_items, doc):
        """
        同一行的对象
        """
        p = doc.add_paragraph()
        self._render_real_items(doc, p, real_items)

    def _render_not_inline(self, real_items, doc):
        """
        不是同行的对象
        """
        for real_item in real_items:
            p = doc.add_paragraph()
            self._render_real_items(p, [real_item])

    def _render_real_items(self, doc, p, real_items):
        """分发处理不同的对象"""
        for real_item in real_items:
            item_type = real_item._TYPE
            if item_type in ["image"]:
                self._render_image_item(p, real_item)
            elif item_type in ["text"]:
                self._render_text_item(p, real_item)
            elif item_type in ["table"]:
                self._render_table_item(doc, real_item)

    def _render_image_item(self, paragraph, image_item):
        """
        处理图片
        paragraph: paragraph对象
        """
        run = paragraph.add_run()
        if image_item.title:
            run.add_text(image_item.title)
        if image_item.identity:
            try:
                pic = self._get_file_stream(image_item.identity)
            except Exception as e:
                print "get_file_stream failed: %s" % e
            else:
                run.add_picture(pic)
        if image_item.text:
            run.add_text(image_item.text)

    def _render_text_item(self, paragraph, text_item):
        """
        处理文本
        """
        run = paragraph.add_run()
        if text_item.text:
            run.add_text(text_item.text)

    def _render_table_item(self, doc, table_item):
        """
        处理不同表格形式
        """
        table_type = table_item.table_type
        tables = {
            1: doc.add_table(1, 4), # 舆论数据模块表格 
            2: doc.add_talbe(1, 4), # 典型报道模块表格
            3: doc.add_talbe(1, 6), # 网民舆论模块表格
            4: doc.add_table(1, 3), #语种分析模块表格 
            5: doc.add_table(1, 3) #港澳台舆论模块表格 
        }

        table_handlers= {
            1: self.__render_1_table,
            2: self.__render_2_table,
            3: self.__render_3_table,
            4: self.__render_4_table,
            5: self.__render_5_table
        }
        if tables.get(table_type) and table_handlers.get(table_type):
            table = table[table_type]
            table_handler = table_handlers[table_type]
            table_handler(table, table_item, table_type)
        else:
            print "get table or table handler wrong, type_type is %s" % table_type

    def __render_table_header(self, table, tb_header_texts=[]):
        tb_header = table.rows[0]
        table.autofit = True
        header_index = 0
        for header_cell in tb_header.cells:
            text = tb_header_texts[header_index]
            cell_p = header_cell.add_paragrph(text)
            header_index += 1

    def __render_1_table(self, table, table_item, table_type=None):
        """table_item.datas数据结构:
        {"data":{"新闻": {"number": "", "groupData":[{"no": "", "name": "", "count": ""}]}}}
            
        """
        tb_header_texts = ["", "序号", "名称", "数量"]
        self.__render_table_header(table, tb_header_texts)
        datas = table_item.datas
        if datas and datas.get('data'):
            data = datas["data"]
            for site_type, value in data.itertiems():
                total = value.get("number", 0)
                group_data = value.get("groupData", [])
                if group_data:
                    group_data_len = len(group_data) 
                    row_count = 0
                    while row_count <= group_data_len:
                        row = table.add_row()
                        row_cells = row.cells
                        # 第二行第一列开始合并上一行第一列
                        if row_count > 0:
                            cell_0 = row_cells[0]
                            cell_0_merge = cell_0.merge(table.cell(row_count-1, 0))
                        if row_count < group_data_len:
                            group_item = group_data[row_count]
                            if group_item:
                                row_cells[1].add_paragraph(group_item.get("no", ""))
                                row_cells[2].add_paragraph(group_item.get("name", ""))
                                row_cells[3].add_paragraph(group_item.get("count", ""))
                        if row_count == group_data_len:
                            row_cells[1].merge(row_cells[2]).add_paragraph(u"小计")
                            row_cells[3].add_paragraph(str(total))
                            cell_0_merge.add_paragraph(site_type)
                        row_count += 1

    def __render_2_table(self, table, table_item, table_type=None):
        table.autofit = True
        datas = table_item.datas
        if datas: 
            data = datas
            count = 0
            seq = 1
            for item in data:
                if item: 
                    title = item.get("title", "")
                    group_number = item["Children"].get("groupNumber", "") if "Children" in item else ""
                    number = item["Children"].get("number", "") if "Children" in item else ""
                    group_name= item["Children"].get("groupName", "") if "Children" in item else ""
                    content = item.get("Content", "")
                    if seq == 1:
                        i = 0
                        while i < 4:
                            if i == 0:
                                table.cell(count, 0).add_paragraph(u"序号")
                                table.cell(count, 1).add_paragraph(seq)
                                table.cell(count, 2).add_paragraph(u"标题")
                                table.cell(count, 3).add_paragraph(title)
                            else:
                                row = table.add_row()
                                count += 1
                                if i == 1:
                                    table.cell(count, 0).add_paragraph(u"参与媒体数量（约）")
                                    table.cell(count, 1).add_paragraph(group_number)
                                    table.cell(count, 2).add_paragraph(u"共发报道数量（约）")
                                    table.cell(count, 3).add_paragraph(number)
                                elif i == 2:
                                    table.cell(count, 0).add_paragraph(u"主要参与媒体")
                                    table.cell(count, 1).merge(table.cell(count, 2)).merge(table.cell(count, 3)).add_paragraph(group_name)
                                elif i == 3:
                                    table.cell(count, 0).add_paragraph(u"内容概要")
                                    table.cell(count, 1).merge(table.cell(count, 2)).merge(table.cell(count, 3)).add_paragraph(content)
                            i += 1
                        
                    else:
                        i = 0
                        while i < 4:
                            row = table.add_row()
                            count += 1
                            if i == 0:
                                table.cell(count, 0).add_paragraph(u"序号")
                                table.cell(count, 1).add_paragraph(seq)
                                table.cell(count, 2).add_paragraph(u"标题")
                                table.cell(count, 3).add_paragraph(title)
                            if i == 1:
                                table.cell(count, 0).add_paragraph(u"参与媒体数量（约）")
                                table.cell(count, 1).add_paragraph(group_number)
                                table.cell(count, 2).add_paragraph(u"共发报道数量（约）")
                                table.cell(count, 3).add_paragraph(number)
                            elif i == 2:
                                table.cell(count, 0).add_paragraph(u"主要参与媒体")
                                table.cell(count, 1).merge(table.cell(count, 2)).merge(table.cell(count, 3)).add_paragraph(group_name)
                            elif i == 3:
                                table.cell(count, 0).add_paragraph(u"内容概要")
                                table.cell(count, 1).merge(table.cell(count, 2)).merge(table.cell(count, 3)).add_paragraph(content)
                            i += 1
                    seq += 1

    def __render_3_table(self, table, table_item, table_type=None):
        tb_header_texts = [u"序号", u"名称", u"微博数量", u"粉丝数量", u"关注数量", u"简介"]
        self.__render_table_header(table, tb_header_texts)
        self.__render_table_body(table, table_item, table_type)
        
    def __render_4_table(self, table, table_item, table_type=None):
        tb_header_texts = [u"序号", u"名称", u"数量"]
        self.__render_table_header(table, tb_header_texts)
        self.__render_table_body(table, table_item, table_type)

    def __render_5_table(self, table, table_item, table_type=None):
        tb_header_texts = [u"序号", u"名称", u"数量"]
        self.__render_table_header(table, tb_header_texts)
        self.__render_table_body(table, table_item, table_type)

    def __render_table_body(self, table, table_item, table_type):
        datas = table_item.datas
        data = []
        if table_type in [3, 4, 5]:
            data = datas
        if data:
            count = 0
            seq = 1
            for item in data:
                row = table.add_row()
                row_cells = row.cells
                if table_type == 3:
                    row_cells[0].add_paragraph(seq)
                    row_cells[1].add_paragraph(item.get("Author", ""))
                    row_cells[2].add_paragraph(item.get("Posts", ""))
                    row_cells[1].add_paragraph(item.get("Fans", ""))
                    row_cells[1].add_paragraph(item.get("Follows", ""))
                    row_cells[1].add_paragraph(item.get("Description", ""))
                elif table_type in [4,]:
                    row_cells[0].add_paragraph(seq)
                    row_cells[1].add_paragraph(item.get("groupName", ""))
                    row_cells[2].add_paragraph(item.get("count", ""))
                seq += 1

    def _get_file_stream(self, file_identity):
        file_manager = self.file_manager
        file_stream = file_manager.get(file_identity)
        return file_stream.read()
